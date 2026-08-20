# -*- coding: utf-8 -*-
"""Word 日报生成：省媒介伊蚊传染病疫情蚊媒监测情况（MM月DD日 20:00）.docx。"""
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

from . import config as C
from . import utils as U


def _fmt_value(v):
    try:
        return f"{float(v):.1f}"
    except (TypeError, ValueError):
        return str(v)


def _point_str(row):
    """监测点字符串：{地市}{区}{街道}{监测地点}（{监测指标值}）。"""
    city, district, street = U.parse_region(row[C.COL_REGION])
    city_full = U.city_with_suffix(city)
    # 东莞、中山：去除“市辖区”
    if U.city_short(city) in ("东莞", "中山") and "市辖区" in district:
        district = ""
    return f"{city_full}{district}{street}{row['监测地点']}（{_fmt_value(row[C.COL_VALUE])}）"


def _region_list_text(df, exclude, raw):
    """按指定顺序排列不重复地市；名称不加“市”（最后一个保留“市”）；排除字段加括号注明。"""
    regions = []
    for r in df[C.COL_REGION]:
        city, _, _ = U.parse_region(r)
        s = U.city_short(city)
        if s and s not in regions:
            regions.append(s)
    ordered = [r for r in C.REGION_ORDER if r in regions]
    ordered += [r for r in regions if r not in C.REGION_ORDER]

    affected = set()
    if exclude:
        mask = raw[C.COL_REGION].astype(str).str.contains(exclude, na=False)
        for r in raw.loc[mask, C.COL_REGION]:
            city, _, _ = U.parse_region(r)
            s = U.city_short(city)
            if s:
                affected.add(s)

    names = []
    for r in ordered:
        names.append(f"{r}（{exclude}除外）" if r in affected else r)

    if not names:
        return ""
    head = "、".join(names[:-1]) if len(names) > 1 else ""
    last = names[-1] + "市"
    return head + ("和" if len(names) > 1 else "") + last


def _overview_text(df, label, target, exclude, raw):
    """总体概述段落。"""
    n = len(df)
    vals = df[C.COL_VALUE].astype(float)
    n_pass = int((vals < C.TH_LOW).sum())
    n_risk = n - n_pass
    pct_pass = n_pass / n * 100 if n else 0.0
    pct_risk = n_risk / n * 100 if n else 0.0

    region_text = _region_list_text(df, exclude, raw)
    regions = []
    for r in df[C.COL_REGION]:
        city, _, _ = U.parse_region(r)
        s = U.city_short(city)
        if s and s not in regions:
            regions.append(s)
    n_regions = len(regions)

    districts = set()
    streets = set()
    for r in df[C.COL_REGION]:
        _, d, s = U.parse_region(r)
        if d:
            districts.add(d)
        if s:
            streets.add(s)

    return (
        f"{target.month}月{target.day}日，对{region_text}共{n_regions}个地市"
        f"{len(districts)}个县区{len(streets)}个镇街{n}个村居/监测点开展了蚊媒密度应急监测。"
        f"监测结果显示：{n_pass}个监测点{label}达标准要求，达标率为{pct_pass:.1f}%"
        f"（{n_pass}/{n}）；{n_risk}个监测点{label}为低中高风险，占{pct_risk:.1f}%"
        f"（{n_risk}/{n}）。"
    )


def _risk_paragraphs(df):
    """高风险、中风险、低风险列表段落（按监测指标值从高到低，数量为 0 时省略）。"""
    paras = []
    for word, key in (("高风险", "高"), ("中风险", "中"), ("低风险", "低")):
        items = []
        for _, row in df.iterrows():
            if U.risk_level_word(row[C.COL_VALUE]) == key:
                items.append((float(row[C.COL_VALUE]), row))
        if not items:
            continue
        items.sort(key=lambda x: x[0], reverse=True)
        pts = [_point_str(r) for _, r in items]
        paras.append(f"{word} {len(items)}个：" + "、".join(pts))
    return paras


def _write_section(doc, df, label, target, exclude, raw):
    if df is None or df.empty:
        doc.add_paragraph(f"本次{label}无监测数据。")
        return
    doc.add_paragraph(_overview_text(df, label, target, exclude, raw))
    for p in _risk_paragraphs(df):
        doc.add_paragraph(p)


def build_word(result, outpath):
    """生成 Word 日报。"""
    doc = Document()
    mm, dd = f"{result.target.month:02d}", f"{result.target.day:02d}"
    title = doc.add_heading(
        f"省媒介伊蚊传染病疫情蚊媒监测情况（{mm}月{dd}日 20:00）", level=0
    )
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_heading("一、布雷图指数（BI）调查评估。", level=1)
    _write_section(doc, result.bi_final, "BI", result.target, result.exclude_field, result.raw)

    doc.add_heading("二、成蚊密度（ADI）快速评估。", level=1)
    _write_section(doc, result.adi_final, "ADI", result.target, result.exclude_field, result.raw)

    doc.save(outpath)
    return outpath
